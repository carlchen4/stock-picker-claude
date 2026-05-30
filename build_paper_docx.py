"""Build the academic paper as a .docx with embedded charts."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Base styling
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

def title(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.bold = True; r.font.size = Pt(15)
def subtitle(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.italic = True; r.font.size = Pt(10.5)
def h1(t): doc.add_heading(t, level=1)
def h2(t): doc.add_heading(t, level=2)
def para(t): doc.add_paragraph(t)
def eq(t):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.font.name = "Cambria Math"; r.font.size = Pt(10.5); p.paragraph_format.left_indent = Inches(0.4)
def figure(path, caption):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(6.0))
    c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = c.add_run(caption); r.italic = True; r.font.size = Pt(9)

title("Sector-Conditional Ensemble Selection with Causal Beta Removal")
subtitle("A Walk-Forward Study of Equity Sub-Portfolios in the Canadian and U.S. Technology Markets")
subtitle("Carl Chen  |  Working Paper, May 2026")

h1("Abstract")
para("We develop and evaluate a systematic cross-sectional equity selection framework that combines per-sector tree ensembles with a partially linear Double Machine Learning (DML) stage for sector-beta removal. The framework is instantiated on two distinct universes: a Canadian large-cap portfolio drawn from the S&P/TSX 60 (Financials, Energy, Industrials, Utilities; benchmark XIU.TO) and a U.S. large-cap technology portfolio (Semiconductors, Cloud/Internet, Hardware; benchmark QQQ). Identical modelling machinery is shared across both markets through a configuration-override architecture, isolating the effect of the investment universe from that of the estimator. Using monthly rebalancing over an 84-month sample (2019-2026) evaluated by a 28-month rolling walk-forward with a one-month embargo, the Canadian model attains an out-of-sample annualised Sharpe ratio of approximately 1.71 and an information ratio (IR) of 0.86, while the U.S. technology model attains a Sharpe of approximately 0.92 with an IR of 1.02 and a Jensen alpha of 13.3% per annum relative to QQQ. We subject both models to a battery of overfitting diagnostics - the Deflated Sharpe Ratio (DSR), the Probability of Backtest Overfitting (PBO), Combinatorial Purged Cross-Validation (CPCV), White's Reality Check (WRC), and Benjamini-Hochberg false-discovery-rate (FDR) control. The Canadian model is rated MODERATE (DSR = 90.5%, PBO = 9.5%); the U.S. model is rated WEAK (DSR ~ 55%, PBO ~ 53%), reflecting a shorter effective history and a single bull-market regime. We further document that, within the U.S. technology universe during 2021-2026, an equally weighted hold-all-names portfolio dominates the selection model on risk-adjusted terms, and report a negative result: linkages between U.S. and Chinese technology equities are economically negligible (|rho| <= 0.31), rendering Chinese ADRs ineffective as either predictive features or portfolio hedges.")
para("Keywords: cross-sectional momentum; ensemble learning; double machine learning; backtest overfitting; deflated Sharpe ratio; sector-neutral portfolio construction.   JEL: G11, G17, C53.")

h1("1. Introduction")
para("A persistent challenge in systematic equity management is the separation of genuine idiosyncratic skill (“alpha”) from passive exposure to common factors (“beta”). When a model is trained to forecast raw forward returns, much of its apparent predictive power may reflect sector or market co-movement that an investor could capture more cheaply through an index fund. This paper addresses that challenge with two design commitments. First, selection is performed within sectors rather than across the full universe, so that the model compares economically similar securities (e.g., one semiconductor firm against another). Second, before the learning stage, sector beta is explicitly removed using a Double Machine Learning partially linear specification, forcing the estimator to learn from the residual idiosyncratic component of returns.")
para("The same estimation pipeline is applied to two markets that differ sharply in character. The Canadian universe is dominated by mature, dividend-paying, lower-volatility sectors. The U.S. technology universe is high-beta, growth-oriented, and - over the sample period - shaped by the artificial-intelligence investment cycle. By holding the methodology fixed and varying only the universe, we isolate the contribution of the investment opportunity set from that of the modelling choices.")
para("Our contribution is threefold: (i) a reproducible two-market evaluation of a sector-conditional ensemble-plus-DML selection model; (ii) a rigorous, multi-test overfitting audit that distinguishes a moderately credible Canadian model from a fragile U.S. model; and (iii) two instructive negative results - the dominance of naive diversification over selection in a melt-up, and the absence of usable cross-market information from Chinese technology equities.")

h1("2. Related Work")
para("The cross-sectional momentum effect documented by Jegadeesh and Titman (1993) motivates our momentum features. Extremely Randomised Trees (Geurts, Ernst, and Wehenkel, 2006) provide a low-variance, low-tuning ensemble well suited to small financial samples. The DML partially linear model (Chernozhukov et al., 2018), building on Robinson (1988), supplies the orthogonalisation used here to remove sector beta. Our overfitting diagnostics follow Bailey and Lopez de Prado (2014) for the Deflated Sharpe Ratio, Lopez de Prado (2018) for CPCV and the PBO, and White (2000) for the Reality Check. Feature-level inference is controlled via the Benjamini-Hochberg (1995) false-discovery-rate procedure.")

h1("3. Data")
para("All price and fundamental data are obtained from a single public source (Yahoo Finance) at monthly frequency, spanning approximately 2019 through 2026 (84 monthly observations). The Canadian universe comprises 30 securities across four sectors plus the XIU.TO benchmark. The U.S. technology universe comprises 21 securities across three sectors plus the QQQ benchmark: Semiconductors (NVDA, AMD, AVGO, QCOM, TXN, AMAT, MU, LRCX, TSM), Cloud/Internet (MSFT, AMZN, META, GOOGL, ANET, PLTR, NET, SNOW), and Hardware (AAPL, TSLA, ARM, INTC). The Software sector was excluded by mandate. Macroeconomic controls include cross-asset series (VIX, gold, oil, the 10-year Treasury yield, an inflation proxy) and, for the U.S. model, the 13-week Treasury bill rate (^IRX) in place of the Bank of Canada overnight rate. Sector ETFs (e.g., SOXX, WCLD, XLK) serve as the treatment series in the DML stage.")
para("Fundamental ratios are aligned to a 45-day point-in-time reporting lag to avoid look-ahead bias. We acknowledge one residual look-ahead approximation: the price-to-sales feature uses the current shares-outstanding figure for all historical dates, because reliable historical share counts were unavailable; the distortion is modest for split-only changes but understates historical valuations for heavy repurchasers.")

h1("4. Methodology")
para("4.1 Feature construction. For each security-month we compute momentum (compressed to two principal components), short-term reversal, realised volatility at 20- and 60-day horizons, the relative strength index, a Bollinger z-score, a 52-week-high ratio, a liquidity rank, and sector-specific macro sensitivities. For the U.S. model we add a trailing-twelve-month price-to-sales ratio.")
para("4.2 Sector-beta removal via DML. Within each sector, forward returns are orthogonalised against the contemporaneous sector-ETF forward return (Eq. 2-3). Estimated treatment coefficients are economically large and significant in every sector (t = 6 to 24).")
para("4.3 Per-sector ensemble. Each sector is modelled independently by an Extremely Randomised Trees estimator deployed as a 50/50 blend of a regressor and a classifier (Eq. 4). Per-sector modelling is load-bearing: collapsing to a single global model degrades the Canadian IR from 1.55 to 0.64.")
para("4.4 Portfolio construction. Securities are ranked by blended score within sector. The portfolio guarantees at least one name per required sector and caps holdings at two per sector, yielding eight names (Canada) or six (U.S.), equally weighted. A stickiness bonus (Eq. 5) reduces turnover; a regime filter reduces holdings under joint high-volatility and below-trend conditions.")
para("4.5 Evaluation protocol. Models are evaluated by a rolling walk-forward with a 28-month training window, a one-month embargo, and a six-month exponential half-life on training weights (Eq. 6). Robustness is assessed by CPCV (15 paths), bootstrap intervals, the DSR (Eq. 10) and PBO, White's Reality Check, and BH-FDR control.")

h2("4.6 Mathematical Formulation")
para("Let r(i,t+1) denote the realised forward return of security i, X(i,t) its feature vector, s the sector of i, and N(t) the cross-section size in month t.")
para("(1) Cross-sectional rank normalisation, mapping each raw feature to an ordinal score in [-1, 1]:")
eq("x̃(i,t) = 2 · rank_t(x(i,t)) / (N(t) + 1) − 1")
para("(2) Partially linear DML model (Robinson 1988; Chernozhukov et al. 2018), with sector-ETF forward return d(s,t+1) as treatment:")
eq("r(i,t+1) = θ_s · d(s,t+1) + g(X(i,t)) + ε(i,t+1),   E[ε | X, d] = 0")
eq("θ̂_s = Σ (d − m̂_d(X))(r − m̂_r(X)) / Σ (d − m̂_d(X))²")
para("where m̂_d, m̂_r are cross-fitted estimates of E[d|X], E[r|X].")
para("(3) De-beta'd learning target:")
eq("r̃(i,t+1) = r(i,t+1) − θ̂_s · d(s,t+1)")
para("(4) Ensemble score (regressor f_reg, classifier f_clf giving P(top-quintile)):")
eq("ŝ(i,t) = ½ · f_reg,s(X(i,t)) + ½ · f_clf,s(X(i,t))")
para("(5) Stickiness adjustment (prior holdings H, bonus b = 0.03):")
eq("s′(i,t) = ŝ(i,t) + b · 1{ i ∈ H(t−1) }")
para("(6) Exponential training weight (half-life h = 6 months):")
eq("w(τ) = 2^( −(t − τ) / h )")
para("(7) Net portfolio return (selected set P, n_new new entries, cost c = 0.001):")
eq("R_p(t+1) = (1/|P|) Σ_{i∈P} r(i,t+1) − c · n_new / |P|")
para("(8) Performance statistics (risk-free = 0; annualisation √12; benchmark R_b):")
eq("Sharpe = √12 · mean(R_p) / std(R_p)")
eq("Sortino = √12 · mean(R_p) / [ √12 · std(R_p · 1{R_p<0}) ]")
eq("IR = √12 · mean(R_p − R_b) / std(R_p − R_b)")
eq("β = Cov(R_p, R_b) / Var(R_b)")
eq("α = mean(R_p) − β · mean(R_b),   α_ann = (1 + α)^12 − 1")
para("(9) Up- and down-capture ratios:")
eq("UC = [Π_{R_b>0}(1+R_p) − 1] / [Π_{R_b>0}(1+R_b) − 1]")
eq("DC = [Π_{R_b<0}(1+R_p) − 1] / [Π_{R_b<0}(1+R_b) − 1]")
para("(10) Deflated Sharpe Ratio (Bailey and Lopez de Prado 2014); Φ the standard normal CDF, SR0 the expected maximum Sharpe over N trials, γ3 skewness, γ4 kurtosis, T observations:")
eq("DSR = Φ( (SR − SR0)√(T−1) / √(1 − γ3·SR + ((γ4−1)/4)·SR²) )")

h1("5. Empirical Results")
para("5.1 Canadian model. Over the out-of-sample window the walk-forward portfolio compounds at an annualised 25.1% against 13.8% for the XIU.TO benchmark, with an annualised Sharpe of approximately 1.7 and an information ratio of 0.86; the CPCV distribution has a mean Sharpe of 1.39 (73% of paths exceeding 1.0), which we regard as the most honest forward estimate. The Canadian wealth path (Figure 1) is markedly smoother than its U.S. counterpart, and its shallower drawdowns (Figure 2) are consistent with the model's higher DSR rating.")
figure("fig5_tsx_equity.png", "Figure 1. Cumulative growth of $1 for the Canadian model versus the XIU.TO benchmark over the out-of-sample walk-forward window.")
figure("fig6_tsx_drawdown.png", "Figure 2. Drawdown (underwater) plot for the Canadian model; shaded area denotes the model, dashed line the benchmark.")
para("5.2 U.S. technology model. Over a 54-month out-of-sample window the model delivers an annualised return of 35.5% against 13.6% for QQQ (excess of 21.8% p.a.), with a Sharpe of 0.92, Sortino of 2.09, and IR of 1.02. Market-sensitivity analysis yields a beta of 1.67, a Jensen alpha of 13.3% p.a., an R-squared of 0.80, an up-capture of 358%, and a down-capture of 116%. The distribution is right-skewed (+0.42). These gains carry materially higher risk: annualised volatility of 38.4% vs 20.9%, maximum drawdown of -45.2% vs -32.6%, and 95% monthly VaR of -13.7% vs -8.9%.")
figure("fig1_equity.png", "Figure 3. Cumulative growth of $1 for the U.S. technology model versus QQQ over the out-of-sample walk-forward window.")
para("The cumulative wealth path (Figure 3) shows the model compounding to roughly 3.9x against 1.8x for QQQ, with larger amplitude in both directions. The underwater plot (Figure 4) confirms deeper troughs, consistent with higher beta.")
figure("fig2_drawdown.png", "Figure 4. Drawdown (underwater) plot, U.S. model; shaded area denotes the model, dashed line the benchmark.")
figure("fig3_dist.png", "Figure 5. Histogram of the U.S. model's monthly returns with mean marked; positive skewness (+0.42).")
para("5.3 Feature attribution. Out-of-sample permutation importance is dominated by liquidity and volatility features and short-term reversal; the price-to-sales feature contributes a small positive increment (delta-IR +0.36). The model's idiosyncratic edge resides predominantly in price- and volume-based reversal signals rather than fundamentally favoured names: several holdings carry negative analyst-implied upside, indicating a contrarian, volatility-driven mechanism.")

h1("6. Robustness and Overfitting Diagnostics")
para("The Canadian model is rated MODERATE: DSR = 90.5%, PBO ~ 9.5%, White's Reality Check at 95.7%. The U.S. model is rated WEAK: DSR ~ 55%, PBO ~ 53%. In both markets, Benjamini-Hochberg FDR control finds zero of 28 features individually significant at the 5% level - an expected consequence of the small effective sample (84 months) rather than evidence of no signal, since the joint portfolio nonetheless generates positive risk-adjusted returns. The divergence is attributable to the U.S. model's shorter usable history (several constituents, notably ARM, list only in 2023) and its single bull-market regime.")

h1("7. Discussion")
para("First, within the U.S. technology universe over 2021-2026, an equally weighted portfolio holding all 21 names dominates the six-name selection model on every risk-adjusted axis (Sharpe 2.14 vs 0.94; maximum drawdown -15% vs -45% on a static basis). Figure 6 contrasts selection against diversification across five construction rules.")
figure("fig4_comparison.png", "Figure 6. Risk-adjusted comparison (Sharpe and Information Ratio) across five portfolio-construction rules, illustrating the dominance of diversification over selection in the 2021-2026 bull regime.")
para("In a broad melt-up, where cross-sectional dispersion is high but persistence is low (best and worst names rotate annually), selection systematically forgoes eventual winners. The selection model's value is therefore regime-contingent: it is expected to earn its keep when avoiding losers matters - in dispersed or declining markets - not in rising-tide conditions. Second, the attractive asymmetry of the U.S. model (up-capture 358% vs down-capture 116%) and its positive Jensen alpha must be read cautiously, as they are estimated over a single favourable regime and accompanied by a WEAK overfitting rating.")

h1("8. A Negative Result: Cross-Market Linkage with Chinese Technology Equities")
para("We tested whether Chinese and Hong Kong technology equities (Alibaba, Tencent, PDD, JD, Baidu, and the KWEB ETF) provide predictive or diversifying value. Contemporaneous monthly correlations are near zero (0.00-0.28); one-month lead correlations are weakly negative (-0.10 to -0.31), indicating no usable predictive content. As an overlay, blending KWEB into a U.S. technology book monotonically reduced the Sharpe ratio (0.80 to 0.54 at a 20% allocation), because KWEB delivered -16.6% annualised and fell alongside U.S. technology in the 2022 drawdown. The markets are economically decoupled - Chinese names are driven by domestic regulatory, delisting, and macro factors - so Chinese ADRs serve as neither a feature nor a hedge.")

h1("9. Limitations and Future Work")
para("The principal limitation is sample size: 84 monthly observations cannot support strong feature-level inference, and both models inhabit fragile local optima in which incremental changes routinely degrade out-of-sample performance. Single-vendor data introduce quality risk, and the price-to-sales feature retains a modest look-ahead approximation. The most credible path to higher confidence is not further in-sample experimentation - which inflates the probability of backtest overfitting - but the accumulation of genuine out-of-sample records. Future work includes regime-adaptive position sizing and a permanent equal-weight control arm.")

h1("10. Conclusion")
para("A shared sector-conditional ensemble framework with causal beta removal produces a moderately credible Canadian model and a high-beta, positive-alpha but statistically fragile U.S. technology model. The broader lesson is methodological humility: rigorous overfitting diagnostics, an honest CPCV forward estimate, and a naive diversification control collectively temper the optimistic picture of a raw backtest. Active selection is a regime-contingent tool, not an unconditional source of excess return.")

h1("References")
for ref in [
 "Bailey, D. H., and M. Lopez de Prado (2014). The Deflated Sharpe Ratio. Journal of Portfolio Management 40(5): 94-107.",
 "Benjamini, Y., and Y. Hochberg (1995). Controlling the False Discovery Rate. JRSS B 57(1): 289-300.",
 "Chernozhukov, V., et al. (2018). Double/Debiased Machine Learning for Treatment and Structural Parameters. Econometrics Journal 21(1): C1-C68.",
 "Geurts, P., D. Ernst, and L. Wehenkel (2006). Extremely Randomized Trees. Machine Learning 63(1): 3-42.",
 "Jegadeesh, N., and S. Titman (1993). Returns to Buying Winners and Selling Losers. Journal of Finance 48(1): 65-91.",
 "Lopez de Prado, M. (2018). Advances in Financial Machine Learning. Wiley.",
 "Robinson, P. M. (1988). Root-N-Consistent Semiparametric Regression. Econometrica 56(4): 931-954.",
 "White, H. (2000). A Reality Check for Data Snooping. Econometrica 68(5): 1097-1126.",
]:
    doc.add_paragraph(ref, style="Normal")

h1("Appendix A. Implementation Notes")
para("The two models share a single ~5,000-line estimation codebase (picker.py); the U.S. model (picker_us.py) is a ~300-line configuration override that substitutes the universe, benchmark, sector definitions, macro controls, and short-rate series, and adds the price-to-sales feature, before invoking the shared pipeline. This architecture guarantees that methodological improvements propagate to both markets and that observed performance differences are attributable to the investment universe rather than to divergent code paths.")

doc.save("picker_paper.docx")
print("SAVED picker_paper.docx")
